import os
import io
from typing import Optional, Dict, Any
from PIL import Image
from PyPDF2 import PdfReader, PdfWriter, PdfMerger
from docx import Document
from docx.shared import Inches
import pandas as pd
from openpyxl import load_workbook
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from flask import current_app
import pytesseract
from pdf2docx import Converter
import tempfile

class FileConverter:
    """Classe para conversão universal de arquivos"""
    
    @staticmethod
    def pdf_to_text(pdf_path: str) -> str:
        """Extrai texto de um PDF"""
        text = ""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            return f"Erro ao extrair texto: {str(e)}"
    
    @staticmethod
    def pdf_to_docx(pdf_path: str, output_path: str = None) -> str:
        """Converte PDF para DOCX"""
        if output_path is None:
            output_path = pdf_path.rsplit('.', 1)[0] + '.docx'
        
        try:
            # Tenta usar pdf2docx primeiro
            cv = Converter(pdf_path)
            cv.convert(output_path)
            cv.close()
            return output_path
        except Exception as e:
            # Fallback method usando PyPDF2 e python-docx
            try:
                text = FileConverter.pdf_to_text(pdf_path)
                doc = Document()
                for paragraph in text.split('\n'):
                    if paragraph.strip():
                        doc.add_paragraph(paragraph)
                doc.save(output_path)
                return output_path
            except Exception as e2:
                raise Exception(f"Erro na conversão PDF para DOCX: {str(e2)}")
    
    @staticmethod
    def docx_to_pdf(docx_path: str, output_path: str = None) -> str:
        """Converte DOCX para PDF"""
        if output_path is None:
            output_path = docx_path.rsplit('.', 1)[0] + '.pdf'
        
        try:
            doc = Document(docx_path)
            
            # Create PDF with reportlab
            pdf_doc = SimpleDocTemplate(
                output_path,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            
            styles = getSampleStyleSheet()
            story = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    story.append(Paragraph(paragraph.text, styles['Normal']))
                    story.append(Spacer(1, 12))
            
            pdf_doc.build(story)
            return output_path
        except Exception as e:
            raise Exception(f"Erro na conversão DOCX para PDF: {str(e)}")
    
    @staticmethod
    def image_to_pdf(image_path: str, output_path: str = None) -> str:
        """Converte imagem para PDF"""
        if output_path is None:
            output_path = image_path.rsplit('.', 1)[0] + '.pdf'
        
        try:
            image = Image.open(image_path)
            if image.mode == 'RGBA':
                image = image.convert('RGB')
            
            image.save(output_path, 'PDF', resolution=100.0, quality=95)
            return output_path
        except Exception as e:
            raise Exception(f"Erro na conversão de imagem para PDF: {str(e)}")
    
    @staticmethod
    def excel_to_pdf(excel_path: str, output_path: str = None) -> str:
        """Converte Excel/CSV para PDF"""
        if output_path is None:
            output_path = excel_path.rsplit('.', 1)[0] + '.pdf'
        
        try:
            # Detecta se é CSV ou Excel
            if excel_path.lower().endswith('.csv'):
                df = pd.read_csv(excel_path)
            else:
                df = pd.read_excel(excel_path)
            
            # Cria PDF com reportlab
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            elements = []
            
            # Adiciona título
            styles = getSampleStyleSheet()
            elements.append(Paragraph("Dados da Planilha", styles['Heading1']))
            elements.append(Spacer(1, 12))
            
            # Converte DataFrame para tabela
            data = [df.columns.tolist()] + df.values.tolist()
            
            # Cria tabela
            table = Table(data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 14),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 10),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            elements.append(table)
            doc.build(elements)
            
            return output_path
        except Exception as e:
            raise Exception(f"Erro na conversão Excel/CSV para PDF: {str(e)}")
    
    @staticmethod
    def merge_pdfs(pdf_files: list, output_path: str) -> str:
        """Mescla múltiplos PDFs em um único arquivo"""
        try:
            merger = PdfMerger()
            
            for pdf in pdf_files:
                merger.append(pdf)
            
            merger.write(output_path)
            merger.close()
            
            return output_path
        except Exception as e:
            raise Exception(f"Erro ao mesclar PDFs: {str(e)}")
    
    @staticmethod
    def split_pdf(pdf_path: str, output_dir: str, pages_per_file: int = 1) -> list:
        """Divide um PDF em múltiplos arquivos"""
        try:
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)
            
            output_files = []
            with open(pdf_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                total_pages = len(pdf_reader.pages)
                
                for start_page in range(0, total_pages, pages_per_file):
                    pdf_writer = PdfWriter()
                    end_page = min(start_page + pages_per_file, total_pages)
                    
                    for page_num in range(start_page, end_page):
                        pdf_writer.add_page(pdf_reader.pages[page_num])
                    
                    output_filename = os.path.join(
                        output_dir,
                        f"split_{start_page + 1}-{end_page}.pdf"
                    )
                    
                    with open(output_filename, 'wb') as output_file:
                        pdf_writer.write(output_file)
                    
                    output_files.append(output_filename)
            
            return output_files
        except Exception as e:
            raise Exception(f"Erro ao dividir PDF: {str(e)}")
    
    @staticmethod
    def compress_pdf(pdf_path: str, output_path: str = None) -> str:
        """Comprime um arquivo PDF"""
        if output_path is None:
            output_path = pdf_path.rsplit('.', 1)[0] + '_compressed.pdf'
        
        try:
            reader = PdfReader(pdf_path)
            writer = PdfWriter()
            
            for page in reader.pages:
                page.compress_content_streams()
                writer.add_page(page)
            
            with open(output_path, 'wb') as output_file:
                writer.write(output_file)
            
            return output_path
        except Exception as e:
            raise Exception(f"Erro ao comprimir PDF: {str(e)}")
    
    @staticmethod
    def add_watermark(pdf_path: str, watermark_text: str, output_path: str = None) -> str:
        """Adiciona marca d'água a um PDF"""
        if output_path is None:
            output_path = pdf_path.rsplit('.', 1)[0] + '_watermarked.pdf'
        
        try:
            # Criar PDF com marca d'água
            watermark_pdf = io.BytesIO()
            can = canvas.Canvas(watermark_pdf, pagesize=letter)
            can.setFont("Helvetica", 50)
            can.setFillColorRGB(0.5, 0.5, 0.5, 0.3)  # Cinza com transparência
            can.saveState()
            can.translate(300, 400)
            can.rotate(45)
            can.drawCentredString(0, 0, watermark_text)
            can.restoreState()
            can.save()
            
            watermark_pdf.seek(0)
            watermark = PdfReader(watermark_pdf)
            watermark_page = watermark.pages[0]
            
            # Aplicar marca d'água
            pdf_reader = PdfReader(pdf_path)
            pdf_writer = PdfWriter()
            
            for page in pdf_reader.pages:
                page.merge_page(watermark_page)
                pdf_writer.add_page(page)
            
            with open(output_path, 'wb') as output_file:
                pdf_writer.write(output_file)
            
            return output_path
        except Exception as e:
            raise Exception(f"Erro ao adicionar marca d'água: {str(e)}")
    
    @staticmethod
    def protect_pdf(pdf_path: str, password: str, output_path: str = None) -> str:
        """Protege um PDF com senha"""
        if output_path is None:
            output_path = pdf_path.rsplit('.', 1)[0] + '_protected.pdf'
        
        try:
            pdf_reader = PdfReader(pdf_path)
            pdf_writer = PdfWriter()
            
            for page in pdf_reader.pages:
                pdf_writer.add_page(page)
            
            pdf_writer.encrypt(password)
            
            with open(output_path, 'wb') as output_file:
                pdf_writer.write(output_file)
            
            return output_path
        except Exception as e:
            raise Exception(f"Erro ao proteger PDF: {str(e)}")
    
    @staticmethod
    def ocr_pdf(pdf_path: str, language: str = 'por+eng') -> str:
        """Executa OCR em um PDF escaneado"""
        try:
            # Primeiro converte PDF para imagens
            from pdf2image import convert_from_path
            
            images = convert_from_path(pdf_path, dpi=300)
            text = ""
            
            for i, image in enumerate(images):
                # Executa OCR em cada página
                page_text = pytesseract.image_to_string(image, lang=language)
                text += f"--- Página {i+1} ---\n{page_text}\n\n"
            
            return text
        except Exception as e:
            raise Exception(f"Erro no OCR: {str(e)}")
    
    @staticmethod
    def rotate_pdf(pdf_path: str, rotation: int, output_path: str = None) -> str:
        """Rotaciona páginas do PDF"""
        if output_path is None:
            output_path = pdf_path.rsplit('.', 1)[0] + f'_rotated_{rotation}.pdf'
        
        try:
            pdf_reader = PdfReader(pdf_path)
            pdf_writer = PdfWriter()
            
            for page in pdf_reader.pages:
                page.rotate(rotation)
                pdf_writer.add_page(page)
            
            with open(output_path, 'wb') as output_file:
                pdf_writer.write(output_file)
            
            return output_path
        except Exception as e:
            raise Exception(f"Erro ao rotacionar PDF: {str(e)}")
    
    @staticmethod
    def extract_pages(pdf_path: str, page_numbers: list, output_path: str = None) -> str:
        """Extrai páginas específicas de um PDF"""
        if output_path is None:
            output_path = pdf_path.rsplit('.', 1)[0] + '_extracted.pdf'
        
        try:
            pdf_reader = PdfReader(pdf_path)
            pdf_writer = PdfWriter()
            
            for page_num in page_numbers:
                if 0 <= page_num < len(pdf_reader.pages):
                    pdf_writer.add_page(pdf_reader.pages[page_num])
            
            with open(output_path, 'wb') as output_file:
                pdf_writer.write(output_file)
            
            return output_path
        except Exception as e:
            raise Exception(f"Erro ao extrair páginas: {str(e)}")
    
    @staticmethod
    def convert_file(input_path: str, output_format: str, **kwargs) -> str:
        """Método genérico para conversão de arquivos"""
        input_ext = input_path.rsplit('.', 1)[1].lower()
        
        # Special handling for PDF to TXT conversion
        if input_ext == 'pdf' and output_format == 'txt':
            text = FileConverter.pdf_to_text(input_path)
            # Create output text file
            output_path = kwargs.get('output_path', input_path.rsplit('.', 1)[0] + '.txt')
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
            return output_path
        
        # Regular conversion map for other formats
        conversion_map = {
            ('pdf', 'docx'): FileConverter.pdf_to_docx,
            ('docx', 'pdf'): FileConverter.docx_to_pdf,
            ('jpg', 'pdf'): FileConverter.image_to_pdf,
            ('jpeg', 'pdf'): FileConverter.image_to_pdf,
            ('png', 'pdf'): FileConverter.image_to_pdf,
            ('xlsx', 'pdf'): FileConverter.excel_to_pdf,
            ('xls', 'pdf'): FileConverter.excel_to_pdf,
            ('csv', 'pdf'): FileConverter.excel_to_pdf,
        }
        
        conversion_func = conversion_map.get((input_ext, output_format))
        
        if conversion_func:
            return conversion_func(input_path, **kwargs)
        else:
            raise ValueError(f"Conversão de {input_ext} para {output_format} não suportada")
    
    @staticmethod
    def get_pdf_info(pdf_path: str) -> Dict[str, Any]:
        """Obtém informações sobre um PDF"""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                info = {
                    'num_pages': len(pdf_reader.pages),
                    'page_size': pdf_reader.pages[0].mediabox if pdf_reader.pages else None,
                    'is_encrypted': pdf_reader.is_encrypted,
                    'metadata': pdf_reader.metadata,
                    'file_size': os.path.getsize(pdf_path)
                }
                return info
        except Exception as e:
            return {'error': str(e)}
    
    @staticmethod
    def extract_images_from_pdf(pdf_path: str, output_dir: str) -> list:
        """Extrai imagens de um PDF"""
        try:
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)
            
            image_files = []
            
            with open(pdf_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    if '/XObject' in page['/Resources']:
                        xObject = page['/Resources']['/XObject'].get_object()
                        
                        for obj in xObject:
                            if xObject[obj]['/Subtype'] == '/Image':
                                size = (xObject[obj]['/Width'], xObject[obj]['/Height'])
                                data = xObject[obj].get_data()
                                
                                if xObject[obj]['/ColorSpace'] == '/DeviceRGB':
                                    mode = "RGB"
                                else:
                                    mode = "P"
                                
                                image_filename = os.path.join(output_dir, f"image_p{page_num}_{obj[1:]}.png")
                                
                                if '/Filter' in xObject[obj]:
                                    if xObject[obj]['/Filter'] == '/FlateDecode':
                                        img = Image.frombytes(mode, size, data)
                                        img.save(image_filename)
                                        image_files.append(image_filename)
            
            return image_files
        except Exception as e:
            raise Exception(f"Erro ao extrair imagens: {str(e)}")
    
    @staticmethod
    def create_pdf_from_text(text: str, output_path: str) -> str:
        """Cria um PDF a partir de texto"""
        try:
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            for paragraph in text.split('\n'):
                if paragraph.strip():
                    story.append(Paragraph(paragraph, styles['Normal']))
                    story.append(Spacer(1, 12))
            
            doc.build(story)
            return output_path
        except Exception as e:
            raise Exception(f"Erro ao criar PDF: {str(e)}")